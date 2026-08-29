"""重写填写《国泰项目周报-20260814.docx》中 二、实验内容 → 1,实验12-金融Agent构建及优化。

结构：概述 → 本周工作内容 → 关键实验结果（表格） → 结果小结。
关键数据（主对比、消融实验）用表格呈现。

用法：python report/fill_weekly_report.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(BASE, "国泰项目周报-20260814.docx")
FONT = "宋体"
BODY_SIZE = 12
TABLE_SIZE = 10.5


def set_run(run, size=BODY_SIZE, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)


def set_para(p, text, bold=False, size=BODY_SIZE, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.alignment = align
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=TABLE_SIZE):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold)


def set_table_borders(tbl):
    """为表格添加全边框（防止模板缺少 Table Grid 样式）。"""
    borders_xml = (
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>' % nsdecls("w")
    )
    tblPr = tbl._tbl.tblPr
    tblPr.append(parse_xml(borders_xml))


def make_table(doc, headers, rows):
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    try:
        tbl.style = "Table Grid"
    except Exception:
        set_table_borders(tbl)
    tbl.autofit = True
    # 表头
    for j, h in enumerate(headers):
        set_cell(tbl.rows[0].cells[j], h, bold=True)
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell(tbl.rows[i + 1].cells[j], val)
    return tbl


CONTENT = [
    # (段落内容, 是否加粗)
    ("本周完成了实验12（金融Agent构建及优化）的全部工作：构建了 Financial Agent 系统并复现 MoCA-Agent 基线，完成主对比、消融、扩展能力实验与定性分析，成果包含 Web 界面、实验报告与汇报 PPT。具体如下：", False),
    ("一、本周工作内容", True),
    ("（1）系统构建：完成包含 Coding Agent、金融文档问答、金融数据分析、数学建模四大模块的 Financial Agent 系统，提供 Streamlit Web 界面；实现安全代码执行沙箱（独立进程、超时隔离、图表自动保存）与 DeepSeek 大语言模型统一封装（记录每次调用的 Token 用量与耗时）。", False),
    ("（2）基线复现：复现 MoCA-Agent（MoCA-Fin）完整流程——Claim 目录构建、专家交易员市场、市场清算、代码综合与验证修复，作为对比基线。", False),
    ("（3）实验框架：整理 FinQA、FinanceMath 公开数据集，固定采样 FinQA 20 题 + FinanceMath 20 题作为主测试集，实现统一输入接口、公平性设置与统一评价脚本。", False),
    ("（4）实验执行与成果：完成主对比实验、两组消融实验、扩展能力实验（Coding 6 个任务、端到端 PDF 问答 8 题、数学建模 4 个任务）及 4 个定性案例分析；生成 Web 界面运行截图、实验报告、汇报 PPT 与操作手册。", False),
    ("二、关键实验结果", True),
    ("表1 主对比实验结果（FinQA 20 题 + FinanceMath 20 题）", True),
    # ↓ 此处插入表1
    ("表2 消融实验结果（固定 24 题）", True),
    # ↓ 此处插入表2
    ("三、结果小结", True),
    ("Financial Agent 整体准确率 77.50%，高于 MoCA-Agent 的 55.00%，平均耗时（4.44s vs 18.08s）与 Token 消耗（1128 vs 7366）也显著更低；消融实验表明，任务专用 Prompt 与代码执行机制均对金融数值推理准确率有正面贡献。", False),
]

TABLE1_HEADERS = ["方法", "FinQA", "FinanceMath", "整体准确率", "代码执行成功率", "任务成功率", "平均耗时", "平均Token"]
TABLE1_ROWS = [
    ["MoCA-Agent", "60.00%", "50.00%", "55.00%", "100%", "100%", "18.08s", "7366"],
    ["Financial Agent", "75.00%", "80.00%", "77.50%", "100%", "97.50%", "4.44s", "1128"],
]

TABLE2_HEADERS = ["方法", "FinQA 准确率", "FinanceMath 准确率", "整体准确率"]
TABLE2_ROWS = [
    ["w/o Task-specific Prompt", "66.67%", "83.33%", "75.00%"],
    ["w/o Code Execution", "66.67%", "83.33%", "75.00%"],
    ["Full Financial Agent", "75.00%", "83.33%", "79.17%"],
]


def main():
    doc = Document(SRC)

    # 定位“1,实验12-金融Agent构建及优化：”标题（带冒号的是实验内容标题，区别于本周工作内容列表项）
    start = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("1,实验12-金融Agent构建及优化") and t.endswith("："):
            start = i
            break
    if start is None:
        print("未找到目标段落。")
        return

    # 找到下一个条目标题（以"2,"开头）作为内容区结束，清空该区间全部段落
    end = len(doc.paragraphs)
    for j in range(start + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].text.strip().startswith("2,"):
            end = j
            break
    for j in range(start + 1, end):
        doc.paragraphs[j].clear()
    area = [doc.paragraphs[j] for j in range(start + 1, end)]

    # 逐个写入正文（CONTENT 中的段落）
    cursor = 0
    for text, bold in CONTENT:
        while cursor < len(area) and area[cursor].text.strip():
            cursor += 1
        set_para(area[cursor], text, bold=bold)
        cursor += 1

    # 定位表1、表2 标题段（作为表格插入锚点）
    anchor1 = None
    anchor2 = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("表1"):
            anchor1 = p
        elif p.text.strip().startswith("表2"):
            anchor2 = p

    # 建表并移动到对应标题之后
    t1 = make_table(doc, TABLE1_HEADERS, TABLE1_ROWS)
    if anchor1 is not None:
        anchor1._p.addnext(t1._tbl)
    t2 = make_table(doc, TABLE2_HEADERS, TABLE2_ROWS)
    if anchor2 is not None:
        anchor2._p.addnext(t2._tbl)

    doc.save(SRC)
    print("已重写并保存:", SRC)


if __name__ == "__main__":
    main()
